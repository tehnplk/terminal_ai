# This module is part of the tools package split from tools/__init__.py.

import os

import docx
import openpyxl
from rich.markup import escape
from rich.panel import Panel
from rich.syntax import Syntax

from . import runtime


def create_docx_file(filename: str, content: list) -> str:
    """
    Creates a Microsoft Word (.docx) document containing paragraphs, headings, lists, and tables.
    
    Args:
        filename: The output filename/path (relative to current working directory).
        content: A list of content blocks. Each block is a dictionary:
                 - Heading: {"type": "heading", "text": "Heading text", "level": 1}
                 - Paragraph: {"type": "paragraph", "text": "Paragraph text"}
                 - Bullet List: {"type": "list_bullet", "text": "Item text"}
                 - Numbered List: {"type": "list_number", "text": "Item text"}
                 - Table: {"type": "table", "table_data": [["Col 1", "Col 2"], ["Val 1", "Val 2"]]}
    """
    try:
        # Resolve full path relative to tracking runtime.cwd
        filepath = runtime.resolve_filepath(filename)
        
        # Verify extension is docx
        ext = os.path.splitext(filepath)[1].lower()
        if ext != '.docx':
            return "Error: Output filename must have a '.docx' extension."
            
        # 1. Print Proposed File Creation Panel
        runtime.console.print()
        preview = f"Word document creation with {len(content)} content blocks."
        runtime.console.print(Panel(
            preview,
            title=f"[bold yellow]📝 Proposed DOCX Creation: {escape(filename)}[/bold yellow]",
            subtitle=f"[dim]CWD: {escape(runtime.cwd)}[/dim]",
            border_style="yellow",
            expand=False
        ))
        
        # 2. Auto-approve since it's non-destructive
        runtime.console.print(f"[yellow]Auto-approving file write (non-destructive): {escape(filename)}[/yellow]")
        
        # Ensure parent directories exist
        os.makedirs(os.path.dirname(filepath), exist_ok=True)
        
        # Generate docx
        doc = docx.Document()
        for idx, block in enumerate(content):
            if not isinstance(block, dict):
                continue
            b_type = block.get("type", "paragraph")
            b_text = block.get("text", "")
            
            if b_type == "heading":
                level = block.get("level", 1)
                doc.add_heading(b_text, level=level)
            elif b_type == "paragraph":
                doc.add_paragraph(b_text)
            elif b_type == "list_bullet":
                doc.add_paragraph(b_text, style='List Bullet')
            elif b_type == "list_number":
                doc.add_paragraph(b_text, style='List Number')
            elif b_type == "table":
                table_data = block.get("table_data", [])
                if table_data and isinstance(table_data, list):
                    rows = len(table_data)
                    cols = len(table_data[0]) if rows > 0 else 0
                    if rows > 0 and cols > 0:
                        table = doc.add_table(rows=rows, cols=cols)
                        for r_idx, row in enumerate(table_data):
                            for c_idx, val in enumerate(row):
                                if c_idx < len(table.columns):
                                    table.cell(r_idx, c_idx).text = str(val)
            else:
                # Fallback to paragraph for safety
                doc.add_paragraph(str(block))
                
        doc.save(filepath)
        
        runtime.console.print(f"[green]✓ Successfully wrote file: {escape(filename)}[/green]")
        return f"Success: Word document '{filename}' was successfully created."
        
    except Exception as e:
        runtime.console.print(f"[red]❌ Error creating docx file: {escape(str(e))}[/red]")
        return f"Error creating docx file '{filename}': {str(e)}"


def create_xlsx_file(filename: str, sheets: dict) -> str:
    """
    Creates a Microsoft Excel (.xlsx) workbook with one or more sheets containing tabular data.
    
    Args:
        filename: The output filename/path (relative to current working directory).
        sheets: A dictionary where keys are sheet names and values are 2D arrays (lists of lists) of data.
    """
    try:
        # Resolve full path relative to tracking runtime.cwd
        filepath = runtime.resolve_filepath(filename)
        
        # Verify extension is xlsx
        ext = os.path.splitext(filepath)[1].lower()
        if ext != '.xlsx':
            return "Error: Output filename must have a '.xlsx' extension."
            
        # 1. Print Proposed File Creation Panel
        runtime.console.print()
        preview = f"Excel workbook creation with {len(sheets)} sheet(s)."
        runtime.console.print(Panel(
            preview,
            title=f"[bold yellow]📝 Proposed XLSX Creation: {escape(filename)}[/bold yellow]",
            subtitle=f"[dim]CWD: {escape(runtime.cwd)}[/dim]",
            border_style="yellow",
            expand=False
        ))
        
        # 2. Auto-approve since it's non-destructive
        runtime.console.print(f"[yellow]Auto-approving file write (non-destructive): {escape(filename)}[/yellow]")
        
        # Ensure parent directories exist
        os.makedirs(os.path.dirname(filepath), exist_ok=True)
        
        # Generate xlsx
        wb = openpyxl.Workbook()
        # Remove default sheet
        default_sheet = wb.active
        wb.remove(default_sheet)
        
        for sheet_name, data in sheets.items():
            ws = wb.create_sheet(title=sheet_name)
            if isinstance(data, list):
                for row in data:
                    if isinstance(row, list):
                        ws.append(row)
                    else:
                        ws.append([row])
            else:
                ws.append([str(data)])
                
        # If no sheets were actually created, add a default one back
        if not wb.sheetnames:
            wb.create_sheet(title="Sheet")
            
        wb.save(filepath)
        
        runtime.console.print(f"[green]✓ Successfully wrote file: {escape(filename)}[/green]")
        return f"Success: Excel workbook '{filename}' was successfully created."
        
    except Exception as e:
        runtime.console.print(f"[red]❌ Error creating xlsx file: {escape(str(e))}[/red]")
        return f"Error creating xlsx file '{filename}': {str(e)}"


def create_html_file(filename: str, html_content: str) -> str:
    """
    Creates or overwrites an HTML file (.html or .htm) in the project directory.
    
    Args:
        filename: The name or path of the HTML file to create (relative to the current working directory, e.g., 'index.html').
        html_content: The HTML content/markup to write to the file.
    """
    try:
        # Resolve full path relative to tracking runtime.cwd
        filepath = runtime.resolve_filepath(filename)
        
        # Verify extension is html/htm
        ext = os.path.splitext(filepath)[1].lower()
        if ext not in ('.html', '.htm'):
            return "Error: Output filename must have a '.html' or '.htm' extension."
            
        # 1. Print Proposed File Creation Panel
        runtime.console.print()
        content_lines = html_content.splitlines()
        preview = "\n".join(content_lines[:10])
        if len(content_lines) > 10:
            preview += f"\n... and {len(content_lines) - 10} more lines ..."
            
        runtime.console.print(Panel(
            Syntax(preview, "html", theme="monokai", line_numbers=True),
            title=f"[bold yellow]📝 Proposed HTML Creation: {escape(filename)}[/bold yellow]",
            subtitle=f"[dim]CWD: {escape(runtime.cwd)}[/dim]",
            border_style="yellow",
            expand=False
        ))
        
        # 2. Auto-approve since it's non-destructive
        runtime.console.print(f"[yellow]Auto-approving file write (non-destructive): {escape(filename)}[/yellow]")
        
        # Ensure parent directories exist
        os.makedirs(os.path.dirname(filepath), exist_ok=True)
        
        # Write content
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(html_content)
            
        runtime.console.print(f"[green]✓ Successfully wrote file: {escape(filename)}[/green]")
        return f"Success: HTML file '{filename}' was successfully created/written."
        
    except Exception as e:
        runtime.console.print(f"[red]❌ Error creating HTML file: {escape(str(e))}[/red]")
        return f"Error creating HTML file '{filename}': {str(e)}"

