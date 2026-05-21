import os
import sys
import subprocess
import re
import json
from datetime import datetime, timezone, timedelta

import docx
import openpyxl
import web_cli
from rich.console import Console
from rich.panel import Panel
from rich.prompt import Confirm
from rich.syntax import Syntax
from rich.markup import escape

console = Console()
cwd = os.getcwd()
auto_approve = False

# Regex to find cd commands
cd_pattern = re.compile(r'(?:^|&&|;)\s*cd\s+("[^"]+"|\'[^\']+\'|[^\s&;]+)', re.IGNORECASE)

def configure(runtime_console=None, initial_cwd=None, approve=False):
    global console, cwd, auto_approve
    if runtime_console is not None:
        console = runtime_console
    if initial_cwd is not None:
        cwd = initial_cwd
    auto_approve = approve

def set_auto_approve(value: bool):
    global auto_approve
    auto_approve = value

def get_cwd() -> str:
    return cwd

def set_cwd(value: str):
    global cwd
    cwd = value

def get_artifact_dir() -> str:
    global cwd
    if getattr(sys, 'frozen', False):
        base_dir = os.path.dirname(os.path.abspath(sys.executable))
    else:
        base_dir = cwd
    return os.path.join(base_dir, "artifact")

def resolve_filepath(filename: str) -> str:
    global cwd
    if os.path.isabs(filename):
        try:
            filename = os.path.relpath(filename, cwd)
        except ValueError:
            filename = os.path.basename(filename)
            
    norm = os.path.normpath(filename)
    parts = []
    for part in norm.split(os.sep):
        if part not in ("..", "", "."):
            parts.append(part)
    filename = os.sep.join(parts)
        
    return os.path.abspath(os.path.join(get_artifact_dir(), filename))

def safe_decode(bytes_data: bytes) -> str:
    """Decodes bytes safely by trying multiple encodings."""
    for encoding in ('utf-8', 'cp1252', 'cp850', 'ansi'):
        try:
            return bytes_data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return bytes_data.decode('utf-8', errors='replace')

def update_cwd_from_command(command: str):
    """Updates the tracking cwd variable if the command contains cd/drive changes."""
    global cwd
    command_clean = command.strip()
    
    # Handle Windows drive changes, e.g. D:
    if re.match(r'^[a-zA-Z]:$', command_clean):
        drive = command_clean.upper()
        if os.path.exists(drive + "\\"):
            cwd = drive + "\\"
            return
            
    # Find all cd matches
    matches = cd_pattern.findall(command)
    for match in matches:
        target = match.strip()
        if (target.startswith('"') and target.endswith('"')) or (target.startswith("'") and target.endswith("'")):
            target = target[1:-1]
        try:
            new_path = os.path.abspath(os.path.join(cwd, target))
            if os.path.exists(new_path) and os.path.isdir(new_path):
                cwd = new_path
        except Exception:
            pass

def is_destructive_command(command: str) -> bool:
    """Checks if a command appears to perform a destructive action (deleting files, folders, or databases)."""
    cmd_lower = command.lower()
    
    # Destructive patterns matching keywords that delete files, folders, or database entities
    destructive_patterns = [
        r'\brm\b',
        r'\bdel\b',
        r'\berase\b',
        r'\brmdir\b',
        r'\brd\b',
        r'\bremove-item\b',
        r'\bunlink\b',
        r'\bdrop\b',
        r'\bdelete\b',
        r'\btruncate\b',
        r'\bshred\b',
        r'\bwipe\b',
        r'\bclean\b',
        r'\bpurge\b',
        r'\bdestroy\b',
    ]
    
    for pattern in destructive_patterns:
        if re.search(pattern, cmd_lower):
            return True
            
    return False

def execute_terminal_command(command: str) -> str:
    """
    Executes a terminal command on the user's local system and returns its stdout and stderr.
    The command will run in the current working directory.
    
    Args:
        command: The terminal command to run.
    """
    global cwd, auto_approve
    
    # 1. Print Proposed Command Panel
    console.print()
    console.print(Panel(
        Syntax(command, "bash", theme="monokai", line_numbers=False),
        title="[bold yellow]⚠️ Proposed Terminal Command[/bold yellow]",
        subtitle=f"[dim]CWD: {escape(cwd)}[/dim]",
        border_style="yellow",
        expand=False
    ))
    
    # 2. Get User Confirmation
    needs_confirmation = not auto_approve
    if needs_confirmation and not is_destructive_command(command):
        needs_confirmation = False
        
    if needs_confirmation:
        confirmed = Confirm.ask(
            "[bold red]⚠️ Destructive operation detected! Do you want to execute this command?[/bold red]",
            default=False
        )
    else:
        if auto_approve:
            console.print("[yellow]Auto-approving command execution (auto_approve mode)...[/yellow]")
        else:
            console.print("[yellow]Auto-approving non-destructive command execution...[/yellow]")
        confirmed = True
        
    if not confirmed:
        console.print("[red]❌ Execution denied by user.[/red]")
        return f"Error: User refused to execute this command: {command}. Suggest an alternative or ask the user for clarification."
        
    # 3. Execute Command
    console.print("[bold green]⏳ Running command...[/bold green]")
    
    # Update our tracking cwd if the command changes directories
    update_cwd_from_command(command)
    
    # Detect interactive commands that may require user input
    interactive_commands = ('date', 'time', 'pause', 'set /p', 'choice', 'more', 'edit', 'nslookup', 'diskpart', 'format', 'chkdsk')
    cmd_lower = command.strip().lower()
    is_interactive = any(cmd_lower == ic or cmd_lower.startswith(ic + ' ') or cmd_lower.startswith(ic + '\n') for ic in interactive_commands)
    
    try:
        if is_interactive:
            # Interactive mode: run with real-time stdin/stdout so user can type input
            console.print("[dim]ℹ️ Interactive command detected — type your input below (or press Ctrl+C to cancel).[/dim]")
            try:
                process = subprocess.Popen(
                    command,
                    shell=True,
                    cwd=cwd,
                    env=os.environ.copy()
                )
                process.wait(timeout=120)
                exit_code = process.returncode
            except KeyboardInterrupt:
                process.kill()
                raise KeyboardInterrupt
            except subprocess.TimeoutExpired:
                process.kill()
                console.print("[red]❌ Interactive command timed out after 120 seconds.[/red]")
                return "Error: Interactive command timed out after 120 seconds."
            
            status_text = "[green]Success (0)[/green]" if exit_code == 0 else f"[red]Failed ({exit_code})[/red]"
            border_color = "green" if exit_code == 0 else "red"
            console.print(Panel(
                "[italic dim]Interactive command finished.[/italic dim]",
                title=f"[bold {border_color}]📋 Command Output ({status_text})[/]",
                border_style=border_color,
                expand=False
            ))
            return f"Exit Code: {exit_code}\nInteractive command completed."
        else:
            # Non-interactive mode: capture stdout/stderr via pipes
            # Run using Popen to allow proper interruption handling
            process = None
            try:
                process = subprocess.Popen(
                    command,
                    shell=True,
                    cwd=cwd,
                    stdout=subprocess.PIPE,
                    stderr=subprocess.PIPE,
                    env=os.environ.copy()
                )
                stdout_bytes, stderr_bytes = process.communicate(timeout=60)
                stdout_str = safe_decode(stdout_bytes)
                stderr_str = safe_decode(stderr_bytes)
                exit_code = process.returncode
            except KeyboardInterrupt:
                if process:
                    try:
                        process.terminate()
                    except Exception:
                        pass
                    try:
                        process.kill()
                    except Exception:
                        pass
                raise KeyboardInterrupt
            except subprocess.TimeoutExpired:
                if process:
                    try:
                        process.terminate()
                    except Exception:
                        pass
                    try:
                        process.kill()
                    except Exception:
                        pass
                console.print("[red]❌ Command timed out after 60 seconds.[/red]")
                return "Error: Command timed out after 60 seconds."
            
            # 4. Display Output Panel
            status_text = "[green]Success (0)[/green]" if exit_code == 0 else f"[red]Failed ({exit_code})[/red]"
            border_color = "green" if exit_code == 0 else "red"
            
            output_content = []
            if stdout_str.strip():
                output_content.append(f"[bold]Standard Output:[/bold]\n{escape(stdout_str.strip())}")
            if stderr_str.strip():
                output_content.append(f"[bold red]Standard Error:[/bold red]\n{escape(stderr_str.strip())}")
            if not stdout_str.strip() and not stderr_str.strip():
                output_content.append("[italic dim]No output received.[/italic dim]")
                
            console.print(Panel(
                "\n\n".join(output_content),
                title=f"[bold {border_color}]📋 Command Output ({status_text})[/]",
                border_style=border_color,
                expand=False
            ))
            
            return f"Exit Code: {exit_code}\n\nSTDOUT:\n{stdout_str}\n\nSTDERR:\n{stderr_str}"
        
    except Exception as e:
        console.print(f"[red]❌ Error running command: {escape(str(e))}[/red]")
        return f"Error running command: {str(e)}"

def read_text_file(filename: str, start_line: int = None, end_line: int = None) -> str:
    """
    Reads the content of a text-based file (like .txt, .md, .py, .toml, etc.) in the project directory.
    Supports reading specific line ranges for large files.
    
    Args:
        filename: The name or path of the file to read (relative to the current working directory).
        start_line: The 1-indexed line number to start reading from (inclusive).
        end_line: The 1-indexed line number to stop reading at (inclusive).
    """
    global cwd
    
    try:
        # Resolve full path relative to tracking cwd
        filepath = resolve_filepath(filename)
        
        # Verify extension is a text file
        ext = os.path.splitext(filepath)[1].lower()
        allowed_extensions = ('.txt', '.csv', '.md', '.py', '.toml', '.json', '.yaml', '.yml', '.ini', '.cfg', '.xml', '.gitignore', '.html', '.htm')
        if ext not in allowed_extensions:
            return f"Error: Only text-based files ({', '.join(allowed_extensions)}) can be read using this tool."
            
        if not os.path.exists(filepath):
            return f"Error: File '{filename}' not found in current directory '{cwd}'."
            
        if os.path.isdir(filepath):
            return f"Error: '{filename}' is a directory, not a file."
            
        with open(filepath, 'r', encoding='utf-8', errors='replace') as f:
            lines = f.readlines()
            
        total_lines = len(lines)
        
        # Handle line range selection
        if start_line is not None or end_line is not None:
            start = (start_line - 1) if start_line is not None else 0
            end = end_line if end_line is not None else total_lines
            
            # Bounds check
            start = max(0, min(start, total_lines))
            end = max(0, min(end, total_lines))
            
            if start >= end:
                return f"Error: Invalid line range {start_line}-{end_line} (total lines: {total_lines})."
                
            selected_lines = lines[start:end]
            content = "".join(selected_lines)
            console.print(f"\n[bold cyan]📖 Reading file: {escape(filename)} (lines {start+1}-{end} of {total_lines})[/bold cyan]")
            return content
            
        # If no range specified, apply a safety limit of 1000 lines
        MAX_AUTO_LINES = 1000
        if total_lines > MAX_AUTO_LINES:
            content = "".join(lines[:MAX_AUTO_LINES])
            console.print(f"\n[bold cyan]📖 Reading file: {escape(filename)} (first {MAX_AUTO_LINES} of {total_lines} lines - TRUNCATED for safety)[/bold cyan]")
            return (
                f"[TRUNCATED - Showing first {MAX_AUTO_LINES} of {total_lines} lines. "
                f"The file is too large to read fully at once. Use the 'start_line' and 'end_line' "
                f"parameters of 'read_text_file' to read specific sections of the file.]\n\n"
                f"{content}"
            )
            
        content = "".join(lines)
        # Display feedback in console
        console.print(f"\n[bold cyan]📖 Reading file: {escape(filename)}[/bold cyan]")
        return content
    except Exception as e:
        return f"Error reading file '{filename}': {str(e)}"

def write_text_file(filename: str, content: str) -> str:
    """
    Creates or overwrites a text-based file (like .txt, .csv, .md, .py, etc.) in the project directory.
    
    Args:
        filename: The name or path of the file to create/write (relative to the current working directory).
        content: The text content to write into the file.
    """
    global cwd, auto_approve
    
    try:
        # Resolve full path relative to tracking cwd
        filepath = resolve_filepath(filename)
        
        # Verify extension is a text-based file
        ext = os.path.splitext(filepath)[1].lower()
        allowed_extensions = ('.txt', '.csv', '.md', '.py', '.toml', '.json', '.yaml', '.yml', '.ini', '.cfg', '.xml', '.gitignore', '.html', '.htm')
        if ext not in allowed_extensions:
            return f"Error: Only text-based files ({', '.join(allowed_extensions)}) can be created using this tool."
            
        # 1. Print Proposed File Creation Panel
        console.print()
        # Preview content (limit to first 10 lines for neat display)
        content_lines = content.splitlines()
        preview = "\n".join(content_lines[:10])
        if len(content_lines) > 10:
            preview += f"\n... and {len(content_lines) - 10} more lines ..."
            
        console.print(Panel(
            Syntax(preview, ext[1:] if ext[1:] else "text", theme="monokai", line_numbers=True),
            title=f"[bold yellow]📝 Proposed File Creation: {escape(filename)}[/bold yellow]",
            subtitle=f"[dim]CWD: {escape(cwd)}[/dim]",
            border_style="yellow",
            expand=False
        ))
        
        # 2. Get User Confirmation
        # Since writing/creating a file is not a deletion, it is non-destructive and auto-approved.
        console.print(f"[yellow]Auto-approving file write (non-destructive): {escape(filename)}[/yellow]")
        confirmed = True
            
        # Ensure parent directories exist
        os.makedirs(os.path.dirname(filepath), exist_ok=True)
        
        # Write content
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(content)
            
        console.print(f"[green]✓ Successfully wrote file: {escape(filename)}[/green]")
        return f"Success: File '{filename}' was successfully created/written."
        
    except Exception as e:
        console.print(f"[red]❌ Error creating file: {escape(str(e))}[/red]")
        return f"Error creating file '{filename}': {str(e)}"

def edit_text_file(filename: str, find_str: str = None, replace_str: str = None, line_number: int = None, content: str = None) -> str:
    """
    Edits a text-based file in the project directory by finding and replacing a text block, 
    replacing a specific line, or appending content.
    
    Args:
        filename: The name or path of the file to edit (relative to current working directory).
        find_str: The exact text block to search for.
        replace_str: The replacement text for find_str.
        line_number: The 1-indexed line number to replace.
        content: The text content to write/replace/append.
    """
    global cwd
    
    try:
        # Resolve full path relative to tracking cwd
        filepath = resolve_filepath(filename)
        
        # Verify extension is a text-based file
        ext = os.path.splitext(filepath)[1].lower()
        allowed_extensions = ('.txt', '.csv', '.md', '.py', '.toml', '.json', '.yaml', '.yml', '.ini', '.cfg', '.xml', '.gitignore', '.html', '.htm')
        if ext not in allowed_extensions:
            return f"Error: Only text-based files ({', '.join(allowed_extensions)}) can be edited using this tool."
            
        if not os.path.exists(filepath):
            return f"Error: File '{filename}' not found in current directory '{cwd}'."
            
        if os.path.isdir(filepath):
            return f"Error: '{filename}' is a directory, not a file."
            
        # 1. Print Proposed File Edit Panel
        console.print()
        edit_desc = ""
        if find_str is not None:
            edit_desc = f"Search and replace in '{escape(filename)}':\n[bold yellow]Find:[/bold yellow]\n{escape(find_str)}\n\n[bold green]Replace:[/bold green]\n{escape(replace_str if replace_str is not None else '')}"
        elif line_number is not None:
            edit_desc = f"Replace line {line_number} in '{escape(filename)}' with:\n{escape(content)}"
        elif content is not None:
            edit_desc = f"Append to '{escape(filename)}':\n{escape(content)}"
        else:
            return "Error: You must provide either find_str, line_number, or content to perform an edit."
            
        console.print(Panel(
            edit_desc,
            title=f"[bold yellow]📝 Proposed File Edit: {escape(filename)}[/bold yellow]",
            subtitle=f"[dim]CWD: {escape(cwd)}[/dim]",
            border_style="yellow",
            expand=False
        ))
        
        # 2. Auto-approve since it's non-destructive
        console.print(f"[yellow]Auto-approving file edit (non-destructive): {escape(filename)}[/yellow]")
        
        # Read current content
        with open(filepath, 'r', encoding='utf-8', errors='replace') as f:
            lines = f.readlines()
            file_content = "".join(lines)
            
        if find_str is not None:
            if find_str not in file_content:
                return f"Error: The target text to find was not found in '{filename}'."
            new_content = file_content.replace(find_str, replace_str if replace_str is not None else "")
        elif line_number is not None:
            idx = line_number - 1
            if idx < 0 or idx >= len(lines):
                return f"Error: Line number {line_number} is out of bounds (total lines in file: {len(lines)})."
            lines[idx] = content + ('\n' if not content.endswith('\n') else '')
            new_content = "".join(lines)
        else:
            new_content = file_content + content
            
        # Write back
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(new_content)
            
        console.print(f"[green]✓ Successfully edited file: {escape(filename)}[/green]")
        return f"Success: File '{filename}' was successfully edited."
        
    except Exception as e:
        console.print(f"[red]❌ Error editing file: {escape(str(e))}[/red]")
        return f"Error editing file '{filename}': {str(e)}"

def create_docx_file(filename: str, content: list) -> str:
    """
    Creates a Microsoft Word (.docx) document containing paragraphs, headings, lists, and tables.
    
    Args:
        filename: The output filename/path (relative to current working directory).
        content: A list of content blocks. Each block is a dictionary:
                 - Heading: {"type": "heading", "text": "Heading text", "level": 1}
                 - Paragraph: {"type": "paragraph", "text": "Paragraph text"}
                 - Bullet List: {"type": "list_bullet", "text": "Item text"}
                 - Numbered List: {"type": "list_number", "text": "Item text"}
                 - Table: {"type": "table", "table_data": [["Col 1", "Col 2"], ["Val 1", "Val 2"]]}
    """
    global cwd
    try:
        # Resolve full path relative to tracking cwd
        filepath = resolve_filepath(filename)
        
        # Verify extension is docx
        ext = os.path.splitext(filepath)[1].lower()
        if ext != '.docx':
            return "Error: Output filename must have a '.docx' extension."
            
        # 1. Print Proposed File Creation Panel
        console.print()
        preview = f"Word document creation with {len(content)} content blocks."
        console.print(Panel(
            preview,
            title=f"[bold yellow]📝 Proposed DOCX Creation: {escape(filename)}[/bold yellow]",
            subtitle=f"[dim]CWD: {escape(cwd)}[/dim]",
            border_style="yellow",
            expand=False
        ))
        
        # 2. Auto-approve since it's non-destructive
        console.print(f"[yellow]Auto-approving file write (non-destructive): {escape(filename)}[/yellow]")
        
        # Ensure parent directories exist
        os.makedirs(os.path.dirname(filepath), exist_ok=True)
        
        # Generate docx
        doc = docx.Document()
        for idx, block in enumerate(content):
            if not isinstance(block, dict):
                continue
            b_type = block.get("type", "paragraph")
            b_text = block.get("text", "")
            
            if b_type == "heading":
                level = block.get("level", 1)
                doc.add_heading(b_text, level=level)
            elif b_type == "paragraph":
                doc.add_paragraph(b_text)
            elif b_type == "list_bullet":
                doc.add_paragraph(b_text, style='List Bullet')
            elif b_type == "list_number":
                doc.add_paragraph(b_text, style='List Number')
            elif b_type == "table":
                table_data = block.get("table_data", [])
                if table_data and isinstance(table_data, list):
                    rows = len(table_data)
                    cols = len(table_data[0]) if rows > 0 else 0
                    if rows > 0 and cols > 0:
                        table = doc.add_table(rows=rows, cols=cols)
                        for r_idx, row in enumerate(table_data):
                            for c_idx, val in enumerate(row):
                                if c_idx < len(table.columns):
                                    table.cell(r_idx, c_idx).text = str(val)
            else:
                # Fallback to paragraph for safety
                doc.add_paragraph(str(block))
                
        doc.save(filepath)
        
        console.print(f"[green]✓ Successfully wrote file: {escape(filename)}[/green]")
        return f"Success: Word document '{filename}' was successfully created."
        
    except Exception as e:
        console.print(f"[red]❌ Error creating docx file: {escape(str(e))}[/red]")
        return f"Error creating docx file '{filename}': {str(e)}"

def create_xlsx_file(filename: str, sheets: dict) -> str:
    """
    Creates a Microsoft Excel (.xlsx) workbook with one or more sheets containing tabular data.
    
    Args:
        filename: The output filename/path (relative to current working directory).
        sheets: A dictionary where keys are sheet names and values are 2D arrays (lists of lists) of data.
    """
    global cwd
    try:
        # Resolve full path relative to tracking cwd
        filepath = resolve_filepath(filename)
        
        # Verify extension is xlsx
        ext = os.path.splitext(filepath)[1].lower()
        if ext != '.xlsx':
            return "Error: Output filename must have a '.xlsx' extension."
            
        # 1. Print Proposed File Creation Panel
        console.print()
        preview = f"Excel workbook creation with {len(sheets)} sheet(s)."
        console.print(Panel(
            preview,
            title=f"[bold yellow]📝 Proposed XLSX Creation: {escape(filename)}[/bold yellow]",
            subtitle=f"[dim]CWD: {escape(cwd)}[/dim]",
            border_style="yellow",
            expand=False
        ))
        
        # 2. Auto-approve since it's non-destructive
        console.print(f"[yellow]Auto-approving file write (non-destructive): {escape(filename)}[/yellow]")
        
        # Ensure parent directories exist
        os.makedirs(os.path.dirname(filepath), exist_ok=True)
        
        # Generate xlsx
        wb = openpyxl.Workbook()
        # Remove default sheet
        default_sheet = wb.active
        wb.remove(default_sheet)
        
        for sheet_name, data in sheets.items():
            ws = wb.create_sheet(title=sheet_name)
            if isinstance(data, list):
                for row in data:
                    if isinstance(row, list):
                        ws.append(row)
                    else:
                        ws.append([row])
            else:
                ws.append([str(data)])
                
        # If no sheets were actually created, add a default one back
        if not wb.sheetnames:
            wb.create_sheet(title="Sheet")
            
        wb.save(filepath)
        
        console.print(f"[green]✓ Successfully wrote file: {escape(filename)}[/green]")
        return f"Success: Excel workbook '{filename}' was successfully created."
        
    except Exception as e:
        console.print(f"[red]❌ Error creating xlsx file: {escape(str(e))}[/red]")
        return f"Error creating xlsx file '{filename}': {str(e)}"

def create_html_file(filename: str, html_content: str) -> str:
    """
    Creates or overwrites an HTML file (.html or .htm) in the project directory.
    
    Args:
        filename: The name or path of the HTML file to create (relative to the current working directory, e.g., 'index.html').
        html_content: The HTML content/markup to write to the file.
    """
    global cwd
    try:
        # Resolve full path relative to tracking cwd
        filepath = resolve_filepath(filename)
        
        # Verify extension is html/htm
        ext = os.path.splitext(filepath)[1].lower()
        if ext not in ('.html', '.htm'):
            return "Error: Output filename must have a '.html' or '.htm' extension."
            
        # 1. Print Proposed File Creation Panel
        console.print()
        content_lines = html_content.splitlines()
        preview = "\n".join(content_lines[:10])
        if len(content_lines) > 10:
            preview += f"\n... and {len(content_lines) - 10} more lines ..."
            
        console.print(Panel(
            Syntax(preview, "html", theme="monokai", line_numbers=True),
            title=f"[bold yellow]📝 Proposed HTML Creation: {escape(filename)}[/bold yellow]",
            subtitle=f"[dim]CWD: {escape(cwd)}[/dim]",
            border_style="yellow",
            expand=False
        ))
        
        # 2. Auto-approve since it's non-destructive
        console.print(f"[yellow]Auto-approving file write (non-destructive): {escape(filename)}[/yellow]")
        
        # Ensure parent directories exist
        os.makedirs(os.path.dirname(filepath), exist_ok=True)
        
        # Write content
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(html_content)
            
        console.print(f"[green]✓ Successfully wrote file: {escape(filename)}[/green]")
        return f"Success: HTML file '{filename}' was successfully created/written."
        
    except Exception as e:
        console.print(f"[red]❌ Error creating HTML file: {escape(str(e))}[/red]")
        return f"Error creating HTML file '{filename}': {str(e)}"


def get_latest_snapshot_content() -> str:
    """Reads the latest .yml snapshot file in .playwright-cli/ and returns its content."""
    dir_path = os.path.join(os.getcwd(), ".playwright-cli")
    if not os.path.exists(dir_path):
        return "Error: No snapshot directory found. You may need to open a page first."
    try:
        files = [f for f in os.listdir(dir_path) if f.startswith("page-") and f.endswith(".yml")]
        if not files:
            return "Error: No snapshots found. The page might not have loaded successfully."
        # Sort alphabetically (chronological by timestamp)
        latest_file = sorted(files)[-1]
        filepath = os.path.join(dir_path, latest_file)
        with open(filepath, 'r', encoding='utf-8', errors='replace') as f:
            content = f.read()
        return f"### Current Page Snapshot ({latest_file}):\n\n```yaml\n{content}\n```"
    except Exception as e:
        return f"Error reading latest snapshot: {str(e)}"

def web_search(query: str) -> str:
    """
    Searches the web using DuckDuckGo via the web_cli helper.
    
    Args:
        query: The search query.
    """
    console.print(f"\n[bold yellow]🔍 Web Search:[/bold yellow] {escape(query)}")
    try:
        return web_cli.search_web(query)
    except Exception as e:
        return f"Error running search: {str(e)}"

def web_fetch(url: str) -> str:
    """
    Fetches the clean text content of a web page using web_cli helper.
    
    Args:
        url: The web page URL.
    """
    console.print(f"\n[bold yellow]🌐 Web Fetch:[/bold yellow] {escape(url)}")
    try:
        return web_cli.fetch_web(url)
    except Exception as e:
        return f"Error running fetch: {str(e)}"

def web_browser_open(url: str) -> str:
    """
    Opens a URL in a browser session using playwright-cli.
    
    Args:
        url: The URL to open.
    """
    console.print(f"\n[bold yellow]🖥️ Opening Browser:[/bold yellow] {escape(url)}")
    try:
        result = subprocess.run(
            ["playwright-cli", "open", url],
            capture_output=True,
            text=True,
            encoding='utf-8',
            errors='replace',
            shell=True,
            cwd=cwd
        )
        if result.returncode != 0:
            return f"Failed to open browser (Exit Code {result.returncode}):\n{result.stderr}\n{result.stdout}"
        return get_latest_snapshot_content()
    except Exception as e:
        return f"Error opening browser: {str(e)}"

def web_browser_action(action: str, target: str = None, text: str = None) -> str:
    """
    Executes an action in the active browser session.
    
    Args:
        action: The action type (e.g. click, fill, type, press, select, hover, reload, go-back, go-forward).
        target: Optional element reference (e.g., e1, e2) or key name (e.g., Enter).
        text: Optional text input value (e.g. for fill/type).
    """
    console.print(f"\n[bold yellow]🖥️ Browser Action:[/bold yellow] {action} target={target} text={text}")
    cmd_args = ["playwright-cli", action]
    if target is not None:
        cmd_args.append(target)
    if text is not None:
        cmd_args.append(text)
    try:
        result = subprocess.run(
            cmd_args,
            capture_output=True,
            text=True,
            encoding='utf-8',
            errors='replace',
            shell=True,
            cwd=cwd
        )
        if result.returncode != 0:
            return f"Action failed (Exit Code {result.returncode}):\n{result.stderr}\n{result.stdout}"
        return get_latest_snapshot_content()
    except Exception as e:
        return f"Error performing browser action: {str(e)}"

def web_browser_close() -> str:
    """Closes the active browser session."""
    console.print("\n[bold yellow]🖥️ Closing Browser[/bold yellow]")
    try:
        result = subprocess.run(
            ["playwright-cli", "close"],
            capture_output=True,
            text=True,
            encoding='utf-8',
            errors='replace',
            shell=True,
            cwd=cwd
        )
        if result.returncode != 0:
            return f"Failed to close browser (Exit Code {result.returncode}):\n{result.stderr}"
        return "Success: Browser session closed successfully."
    except Exception as e:
        return f"Error closing browser: {str(e)}"

def get_current_time() -> str:
    """
    Returns the current date and time in the Asia/Bangkok timezone (UTC+7).
    """
    tz = timezone(timedelta(hours=7))
    now = datetime.now(tz)
    return f"Current date and time in Asia/Bangkok: {now.strftime('%Y-%m-%d %H:%M:%S %Z')}"

# Available functions for the agent loop mapping
available_functions = {
    "execute_terminal_command": execute_terminal_command,
    "read_text_file": read_text_file,
    "write_text_file": write_text_file,
    "edit_text_file": edit_text_file,
    "web_search": web_search,
    "web_fetch": web_fetch,
    "web_browser_open": web_browser_open,
    "web_browser_action": web_browser_action,
    "web_browser_close": web_browser_close,
    "create_docx_file": create_docx_file,
    "create_xlsx_file": create_xlsx_file,
    "create_html_file": create_html_file,
    "get_current_time": get_current_time,
}

# OpenAI/OpenRouter compatible tools schema definition
tools_schema = [
    {
        "type": "function",
        "function": {
            "name": "execute_terminal_command",
            "description": "Executes a terminal command on the user's local system and returns its stdout and stderr. The command will run in the current working directory.",
            "parameters": {
                "type": "object",
                "properties": {
                    "command": {
                        "type": "string",
                        "description": "The terminal command to run."
                    }
                },
                "required": ["command"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "read_text_file",
            "description": "Reads the content of a text-based file (like .txt, .md, .py, .toml, etc.) in the project directory. Supports reading specific line ranges for large files.",
            "parameters": {
                "type": "object",
                "properties": {
                    "filename": {
                        "type": "string",
                        "description": "The name or path of the file to read (relative to the current working directory)."
                    },
                    "start_line": {
                        "type": "integer",
                        "description": "The 1-indexed line number to start reading from (inclusive)."
                    },
                    "end_line": {
                        "type": "integer",
                        "description": "The 1-indexed line number to stop reading at (inclusive)."
                    }
                },
                "required": ["filename"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "write_text_file",
            "description": "Creates or overwrites a text-based file (like .txt, .csv, .md, .py, etc.) in the project directory.",
            "parameters": {
                "type": "object",
                "properties": {
                    "filename": {
                        "type": "string",
                        "description": "The name or path of the file to create/write (relative to the current working directory)."
                    },
                    "content": {
                        "type": "string",
                        "description": "The text content to write into the file."
                    }
                },
                "required": ["filename", "content"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "edit_text_file",
            "description": "Edits an existing text-based file in the project directory by finding and replacing a text block, replacing a specific line, or appending content.",
            "parameters": {
                "type": "object",
                "properties": {
                    "filename": {
                        "type": "string",
                        "description": "The name or path of the file to edit (relative to the current working directory)."
                    },
                    "find_str": {
                        "type": "string",
                        "description": "The exact text block to search for and replace."
                    },
                    "replace_str": {
                        "type": "string",
                        "description": "The replacement text for find_str."
                    },
                    "line_number": {
                        "type": "integer",
                        "description": "The 1-indexed line number to replace."
                    },
                    "content": {
                        "type": "string",
                        "description": "New line content (required when using line_number) or content to append."
                    }
                },
                "required": ["filename"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "web_search",
            "description": "Searches the web for a given query and returns search results (titles, URLs, snippets). Uses DuckDuckGo.",
            "parameters": {
                "type": "object",
                "properties": {
                    "query": {
                        "type": "string",
                        "description": "The search query string."
                    }
                },
                "required": ["query"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "web_fetch",
            "description": "Fetches clean, readable text content from a URL.",
            "parameters": {
                "type": "object",
                "properties": {
                    "url": {
                        "type": "string",
                        "description": "The URL to fetch."
                    }
                },
                "required": ["url"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "web_browser_open",
            "description": "Opens a URL in a browser session and returns the structured page snapshot.",
            "parameters": {
                "type": "object",
                "properties": {
                    "url": {
                        "type": "string",
                        "description": "The URL to open."
                    }
                },
                "required": ["url"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "web_browser_action",
            "description": "Performs an interactive action (click, fill, type, press, select, hover, reload, go-back, go-forward) in the active browser session and returns the updated snapshot.",
            "parameters": {
                "type": "object",
                "properties": {
                    "action": {
                        "type": "string",
                        "description": "The action type (click, fill, type, press, select, hover, reload, go-back, go-forward)."
                    },
                    "target": {
                        "type": "string",
                        "description": "Optional element reference (e.g. e1, e2) or key name (e.g. Enter)."
                    },
                    "text": {
                        "type": "string",
                        "description": "Optional text input value (required for fill/type)."
                    }
                },
                "required": ["action"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "web_browser_close",
            "description": "Closes the active browser session.",
            "parameters": {
                "type": "object",
                "properties": {}
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "create_docx_file",
            "description": "Creates a Microsoft Word (.docx) document containing headings, paragraphs, lists, and tables.",
            "parameters": {
                "type": "object",
                "properties": {
                    "filename": {
                        "type": "string",
                        "description": "The output filename or path (relative to the current working directory, e.g., 'report.docx')."
                    },
                    "content": {
                        "type": "array",
                        "description": "A list of content blocks to add to the document.",
                        "items": {
                            "type": "object",
                            "properties": {
                                "type": {
                                    "type": "string",
                                    "enum": ["heading", "paragraph", "list_bullet", "list_number", "table"],
                                    "description": "The type of content block."
                                },
                                "text": {
                                    "type": "string",
                                    "description": "The text content (for headings, paragraphs, and list items)."
                                },
                                "level": {
                                    "type": "integer",
                                    "description": "The heading level (required/used only when type is 'heading', e.g., 1 for title/main heading, 2 for subheading)."
                                },
                                "table_data": {
                                    "type": "array",
                                    "description": "A 2D array representing table cells (required/used only when type is 'table').",
                                    "items": {
                                        "type": "array"
                                    }
                                }
                            },
                            "required": ["type"]
                        }
                    }
                },
                "required": ["filename", "content"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "create_xlsx_file",
            "description": "Creates a Microsoft Excel (.xlsx) workbook with one or more sheets containing tabular data.",
            "parameters": {
                "type": "object",
                "properties": {
                    "filename": {
                        "type": "string",
                        "description": "The output filename or path (relative to the current working directory, e.g., 'sales.xlsx')."
                    },
                    "sheets": {
                        "type": "object",
                        "description": "A dictionary mapping sheet names to a 2D array of cell values. Example: {'Sales': [['Product', 'Price'], ['Apple', 1.50], ['Banana', 2.00]]}"
                    }
                },
                "required": ["filename", "sheets"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "create_html_file",
            "description": "Creates or overwrites an HTML file (.html or .htm) containing HTML markup.",
            "parameters": {
                "type": "object",
                "properties": {
                    "filename": {
                        "type": "string",
                        "description": "The output filename or path (relative to the current working directory, e.g., 'index.html')."
                    },
                    "html_content": {
                        "type": "string",
                        "description": "The complete HTML markup/content to write into the file."
                    }
                },
                "required": ["filename", "html_content"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "get_current_time",
            "description": "Returns the current date and time in the Asia/Bangkok timezone (UTC+7).",
            "parameters": {
                "type": "object",
                "properties": {}
            }
        }
    }
]

